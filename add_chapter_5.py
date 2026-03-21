#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
添加第五章：系统实现

本脚本用于向毕业论文 Word 文档中添加第五章内容。
包括：前端实现、后端实现、数据处理脚本、关键技术实现等部分。
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_paragraph_formatted(doc, text, font_name='宋体', font_size=12,
                           line_spacing=1.5, indent=0):
    """
    添加格式化段落

    参数：
    - doc: Document 对象
    - text: 段落文本
    - font_name: 字体名称（默认宋体）
    - font_size: 字体大小（默认12pt）
    - line_spacing: 行距（默认1.5倍）
    - indent: 首行缩进（单位：厘米）

    返回：
    - 创建的段落对象
    """
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 左对齐
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)  # 段前间距
    p.paragraph_format.space_after = Pt(0)   # 段后间距
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)  # 首行缩进

    # 设置段落中所有文本的字体
    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)

    return p


def add_heading_formatted(doc, text, level=1):
    """
    添加格式化标题

    参数：
    - doc: Document 对象
    - text: 标题文本
    - level: 标题级别（1=一级标题，2=二级标题，3=三级标题）

    返回：
    - 创建的标题对象
    """
    if level == 1:
        # 一级标题：黑体，16pt，加粗
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(16)
            run.font.bold = True
    elif level == 2:
        # 二级标题：黑体，14pt，加粗
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(14)
            run.font.bold = True
    else:
        # 三级标题：宋体，12pt
        heading = doc.add_heading(text, level=3)
        for run in heading.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)

    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 左对齐
    return heading


def add_chapter_5(doc):
    """
    添加第五章：系统实现

    包含以下内容：
    1. 前端实现（页面结构、关键组件、状态管理、性能优化）
    2. 后端实现（云函数设计、数据库操作、API 设计）
    3. 数据处理脚本（真题数据处理流程、数据质量控制、数据规模）
    4. 关键技术实现（跨平台适配、离线支持、错误处理）
    """

    # 第五章标题
    add_heading_formatted(doc, '5 系统实现', level=1)

    # 5.1 前端实现
    add_heading_formatted(doc, '5.1 前端实现', level=2)

    section_5_1 = """前端采用uni-app + Vue 3 + Vite进行开发，实现了跨平台的用户界面和交互逻辑。

5.1.1 页面结构

系统包含以下主要页面：

（1）首页（index.vue）
- 显示单词列表，支持搜索、筛选、排序等功能
- 采用虚拟滚动技术实现大列表高效渲染
- 支持按难度、重要程度、学习进度等条件筛选
- 支持按创建时间、难度、稳定性等条件排序

（2）单词详情页（word-detail.vue）
- 显示单词的详细信息，包括释义、例句、真题统计等
- 支持编辑单词信息
- 支持查看真题出处和例句
- 支持AI生成例句和标签推荐

（3）复习页面（review.vue）
- 实现FSRS算法的复习流程
- 支持多种题型：选择题、填空题、形近词辨析等
- 显示复习进度和统计信息
- 支持自定义复习配置

（4）统计页面（stats.vue）
- 显示学习统计信息，包括学习进度、复习频率、错误率等
- 支持按时间段统计，如按天、周、月统计
- 显示真题分析信息，包括单词出现频率、题型分布等
- 提供个性化学习建议

（5）错题本（mistakes.vue）
- 显示用户答错的单词
- 支持重点复习错题
- 支持删除错题记录

（6）已掌握单词（mastered-words.vue）
- 显示用户已掌握的单词
- 支持查看掌握时间和复习次数

（7）单词本列表（wordbook-list.vue）
- 显示用户创建的所有单词本
- 支持创建、编辑、删除单词本
- 支持导入CSV文件

（8）快速添加（quick-add.vue）
- 支持快速添加单词
- 支持批量导入

（9）个人中心（my.vue）
- 显示用户信息
- 支持修改个人信息
- 支持数据备份和恢复
- 支持账号注销

（10）登录页面（login.vue）
- 支持账号登录和注册
- 支持密码重置

5.1.2 关键组件

（1）VirtualScroller虚拟滚动组件
- 原理：只渲染可见区域的元素，不渲染不可见的元素
- 性能优化：减少DOM节点数量，降低内存占用
- 实现细节：
  * 监听滚动事件
  * 计算可见范围
  * 动态更新渲染的元素
  * 使用transform实现高效的位置更新

（2）ReviewCard复习卡片组件
- 显示单词和复习选项
- 支持多种题型
- 实时反馈答题结果

（3）StatChart统计图表组件
- 使用ECharts库绘制图表
- 支持多种图表类型：柱状图、折线图、饼图等
- 支持交互式数据探索

5.1.3 状态管理

系统采用Pinia进行状态管理：

（1）全局状态
- 用户信息（用户ID、用户名、头像等）
- 当前单词本
- 应用配置（主题、语言等）

（2）本地状态
- 页面级别的状态（如搜索关键词、筛选条件等）
- 组件级别的状态（如表单输入等）

（3）缓存管理
- 单词缓存
- 复习记录缓存
- 统计数据缓存

5.1.4 性能优化

（1）代码分割
- 按路由分割代码
- 按功能分割代码
- 减少初始包大小

（2）懒加载
- 图片懒加载
- 组件懒加载
- 数据懒加载

（3）缓存策略
- 内存缓存
- 存储缓存
- HTTP缓存

（4）虚拟滚动
- 只渲染可见区域
- 支持大列表渲染
- 提高滚动帧率"""

    add_paragraph_formatted(doc, section_5_1)

    # 5.2 后端实现
    add_heading_formatted(doc, '5.2 后端实现', level=2)

    section_5_2 = """后端采用uniCloud云函数进行开发，实现了用户认证、数据同步等功能。

5.2.1 云函数设计

系统包含以下主要云函数：

（1）user-center云函数
- 处理用户认证和管理
- 功能包括：
  * 用户注册：验证邮箱/手机号，创建用户账号
  * 用户登录：验证用户名和密码，返回Token
  * 密码验证：使用SHA-256加密验证密码
  * 用户信息更新：修改用户昵称、头像等
  * 用户注销：删除用户账号和数据

（2）word-sync云函数
- 处理单词数据的同步和备份
- 功能包括：
  * 上传本地修改：将本地修改的单词上传到云端
  * 下载云端数据：下载云端的最新数据
  * 冲突解决：处理本地和云端数据的冲突
  * 数据备份：备份用户数据到云端
  * 数据恢复：从云端恢复用户数据

（3）user-progress-sync云函数
- 处理学习进度的同步
- 功能包括：
  * 复习记录同步：同步用户的复习记录
  * 统计数据更新：更新用户的学习统计数据

5.2.2 数据库操作

（1）事务处理
- 复习记录的更新采用事务处理
- 确保数据一致性
- 防止数据丢失

（2）并发控制
- 采用乐观锁机制处理并发更新
- 避免数据冲突
- 提高系统吞吐量

（3）数据验证
- 在数据库层面进行数据验证
- 确保数据的有效性
- 防止非法数据写入

5.2.3 API设计

（1）RESTful风格
- 使用HTTP方法表示操作：GET、POST、PUT、DELETE
- 使用URL表示资源：/users、/words、/reviews等
- 使用HTTP状态码表示结果：200、400、401、500等

（2）请求/响应格式
- 请求格式：JSON
- 响应格式：JSON
- 统一的响应结构：{ code, message, data }

（3）错误处理
- 定义统一的错误码
- 提供清晰的错误信息
- 记录详细的错误日志"""

    add_paragraph_formatted(doc, section_5_2)

    # 5.3 数据处理脚本
    add_heading_formatted(doc, '5.3 数据处理脚本', level=2)

    section_5_3 = """数据处理脚本采用Python编写，负责真题数据的处理和预生成。

5.3.1 真题数据处理流程

（1）pdf_to_txt.py：PDF转TXT
- 使用PyMuPDF库提取PDF文本
- 支持中英文混合文本
- 输出纯文本格式

（2）word_stats_from_trueti.py：统计真题词汇
- 对文本进行分词和清洗
- 提取英文单词
- 统计词频、题型分布、年份分布等
- 输出统计结果

（3）build_master_db.py：生成统一主库
- 创建SQLite数据库
- 导入词汇信息
- 创建索引
- 优化查询性能

（4）build_wordbooks.py：生成单词本
- 根据统计数据生成单词本
- 支持按难度、频率等条件分类
- 输出CSV格式

5.3.2 数据质量控制

（1）验证
- 检查单词格式
- 检查中文释义
- 检查例句质量
- 检查标签准确性

（2）清洗
- 去除重复
- 去除无效数据
- 标准化格式
- 修复编码问题

（3）异常处理
- 记录处理过程中的异常
- 提供详细的错误信息
- 支持断点续传

5.3.3 数据规模

真题数据统计：
- 年份范围：1998-2025（28年）
- 题型：完形、阅读、新题型、翻译、写作
- 单词总数：约5000个
- 真题例句：约10000条"""

    add_paragraph_formatted(doc, section_5_3)

    # 5.4 关键技术实现
    add_heading_formatted(doc, '5.4 关键技术实现', level=2)

    section_5_4 = """系统采用了多项关键技术来实现高性能和良好的用户体验。

5.4.1 跨平台适配

（1）平台检测
- 使用uni.getSystemInfoSync()获取系统信息
- 检测平台类型：iOS、Android、Web
- 检测操作系统版本和屏幕尺寸

（2）条件编译
- 使用#ifdef和#endif进行条件编译
- 为不同平台编译不同的代码
- 减少包体积

（3）API兼容性处理
- SQLite（App端）：使用uni.openDatabase()
- localStorage（H5端）：使用uni.setStorageSync()
- 统一的数据库接口

5.4.2 离线支持

（1）本地数据库
- 所有数据都存储在本地数据库
- 支持离线查询和修改

（2）离线缓存
- 缓存频繁访问的数据
- 减少网络请求

（3）同步机制
- 网络恢复后自动同步
- 支持增量同步
- 冲突自动解决

5.4.3 错误处理

（1）全局错误捕获
- 捕获所有未处理的异常
- 记录错误日志
- 提供用户友好的错误提示

（2）用户友好的错误提示
- 避免技术性错误信息
- 提供解决方案建议
- 支持错误反馈

（3）日志系统
- 记录详细的操作日志
- 支持日志级别控制
- 支持日志导出和分析"""

    add_paragraph_formatted(doc, section_5_4)

    doc.add_page_break()


def main():
    """主函数：打开文档、添加第五章、保存文档"""
    # 打开已生成的文档
    doc = Document(r'e:\vocal\wordbook_new\毕业论文_基于FSRS算法和真题数据融合的英语单词学习应用.docx')

    # 添加第五章
    print("添加第五章：系统实现...")
    add_chapter_5(doc)

    # 保存文档
    output_path = r'e:\vocal\wordbook_new\毕业论文_基于FSRS算法和真题数据融合的英语单词学习应用.docx'
    doc.save(output_path)
    print(f"第五章已添加，文档已保存：{output_path}")


if __name__ == '__main__':
    main()

    # 5.1 前端实现
    add_heading_formatted(doc, '5.1 前端实现', level=2)

    section_5_1 = """前端采用uni-app + Vue 3 + Vite进行开发，实现了跨平台的用户界面和交互逻辑。

5.1.1 页面结构

系统包含以下主要页面：

（1）首页（index.vue）
- 显示单词列表，支持搜索、筛选、排序等功能
- 采用虚拟滚动技术实现大列表高效渲染
- 支持按难度、重要程度、学习进度等条件筛选
- 支持按创建时间、难度、稳定性等条件排序

（2）单词详情页（word-detail.vue）
- 显示单词的详细信息，包括释义、例句、真题统计等
- 支持编辑单词信息
- 支持查看真题出处和例句
- 支持AI生成例句和标签推荐

（3）复习页面（review.vue）
- 实现FSRS算法的复习流程
- 支持多种题型：选择题、填空题、形近词辨析等
- 显示复习进度和统计信息
- 支持自定义复习配置

（4）统计页面（stats.vue）
- 显示学习统计信息，包括学习进度、复习频率、错误率等
- 支持按时间段统计，如按天、周、月统计
- 显示真题分析信息，包括单词出现频率、题型分布等
- 提供个性化学习建议

（5）错题本（mistakes.vue）
- 显示用户答错的单词
- 支持重点复习错题
- 支持删除错题记录

（6）已掌握单词（mastered-words.vue）
- 显示用户已掌握的单词
- 支持查看掌握时间和复习次数

（7）单词本列表（wordbook-list.vue）
- 显示用户创建的所有单词本
- 支持创建、编辑、删除单词本
- 支持导入CSV文件

（8）快速添加（quick-add.vue）
- 支持快速添加单词
- 支持批量导入

（9）个人中心（my.vue）
- 显示用户信息
- 支持修改个人信息
- 支持数据备份和恢复
- 支持账号注销

（10）登录页面（login.vue）
- 支持账号登录和注册
- 支持密码重置

5.1.2 关键组件

（1）VirtualScroller虚拟滚动组件
- 原理：只渲染可见区域的元素，不渲染不可见的元素
- 性能优化：减少DOM节点数量，降低内存占用
- 实现细节：
  * 监听滚动事件
  * 计算可见范围
  * 动态更新渲染的元素
  * 使用transform实现高效的位置更新

（2）ReviewCard复习卡片组件
- 显示单词和复习选项
- 支持多种题型
- 实时反馈答题结果

（3）StatChart统计图表组件
- 使用ECharts库绘制图表
- 支持多种图表类型：柱状图、折线图、饼图等
- 支持交互式数据探索

5.1.3 状态管理

系统采用Pinia进行状态管理：

（1）全局状态
- 用户信息（用户ID、用户名、头像等）
- 当前单词本
- 应用配置（主题、语言等）

（2）本地状态
- 页面级别的状态（如搜索关键词、筛选条件等）
- 组件级别的状态（如表单输入等）

（3）缓存管理
- 单词缓存
- 复习记录缓存
- 统计数据缓存

5.1.4 性能优化

（1）代码分割
- 按路由分割代码
- 按功能分割代码
- 减少初始包大小

（2）懒加载
- 图片懒加载
- 组件懒加载
- 数据懒加载

（3）缓存策略
- 内存缓存
- 存储缓存
- HTTP缓存

（4）虚拟滚动
- 只渲染可见区域
- 支持大列表渲染
- 提高滚动帧率"""

    add_paragraph_formatted(doc, section_5_1)

    # 5.2 后端实现
    add_heading_formatted(doc, '5.2 后端实现', level=2)

    section_5_2 = """后端采用uniCloud云函数进行开发，实现了用户认证、数据同步等功能。

5.2.1 云函数设计

系统包含以下主要云函数：

（1）user-center云函数
- 处理用户认证和管理
- 功能包括：
  * 用户注册：验证邮箱/手机号，创建用户账号
  * 用户登录：验证用户名和密码，返回Token
  * 密码验证：使用SHA-256加密验证密码
  * 用户信息更新：修改用户昵称、头像等
  * 用户注销：删除用户账号和数据

（2）word-sync云函数
- 处理单词数据的同步和备份
- 功能包括：
  * 上传本地修改：将本地修改的单词上传到云端
  * 下载云端数据：下载云端的最新数据
  * 冲突解决：处理本地和云端数据的冲突
  * 数据备份：备份用户数据到云端
  * 数据恢复：从云端恢复用户数据

（3）user-progress-sync云函数
- 处理学习进度的同步
- 功能包括：
  * 复习记录同步：同步用户的复习记录
  * 统计数据更新：更新用户的学习统计数据

5.2.2 数据库操作

（1）事务处理
- 复习记录的更新采用事务处理
- 确保数据一致性
- 防止数据丢失

（2）并发控制
- 采用乐观锁机制处理并发更新
- 避免数据冲突
- 提高系统吞吐量

（3）数据验证
- 在数据库层面进行数据验证
- 确保数据的有效性
- 防止非法数据写入

5.2.3 API设计

（1）RESTful风格
- 使用HTTP方法表示操作：GET、POST、PUT、DELETE
- 使用URL表示资源：/users、/words、/reviews等
- 使用HTTP状态码表示结果：200、400、401、500等

（2）请求/响应格式
- 请求格式：JSON
- 响应格式：JSON
- 统一的响应结构：{ code, message, data }

（3）错误处理
- 定义统一的错误码
- 提供清晰的错误信息
- 记录详细的错误日志"""

    add_paragraph_formatted(doc, section_5_2)

    # 5.3 数据处理脚本
    add_heading_formatted(doc, '5.3 数据处理脚本', level=2)

    section_5_3 = """数据处理脚本采用Python编写，负责真题数据的处理和预生成。

5.3.1 真题数据处理流程

（1）pdf_to_txt.py：PDF转TXT
- 使用PyMuPDF库提取PDF文本
- 支持中英文混合文本
- 输出纯文本格式

（2）word_stats_from_trueti.py：统计真题词汇
- 对文本进行分词和清洗
- 提取英文单词
- 统计词频、题型分布、年份分布等
- 输出统计结果

（3）build_master_db.py：生成统一主库
- 创建SQLite数据库
- 导入词汇信息
- 创建索引
- 优化查询性能

（4）build_wordbooks.py：生成单词本
- 根据统计数据生成单词本
- 支持按难度、频率等条件分类
- 输出CSV格式

5.3.2 数据质量控制

（1）验证
- 检查单词格式
- 检查中文释义
- 检查例句质量
- 检查标签准确性

（2）清洗
- 去除重复
- 去除无效数据
- 标准化格式
- 修复编码问题

（3）异常处理
- 记录处理过程中的异常
- 提供详细的错误信息
- 支持断点续传

5.3.3 数据规模

真题数据统计：
- 年份范围：1998-2025（28年）
- 题型：完形、阅读、新题型、翻译、写作
- 单词总数：约5000个
- 真题例句：约10000条"""

    add_paragraph_formatted(doc, section_5_3)

    # 5.4 关键技术实现
    add_heading_formatted(doc, '5.4 关键技术实现', level=2)

    section_5_4 = """系统采用了多项关键技术来实现高性能和良好的用户体验。

5.4.1 跨平台适配

（1）平台检测
- 使用uni.getSystemInfoSync()获取系统信息
- 检测平台类型：iOS、Android、Web
- 检测操作系统版本和屏幕尺寸

（2）条件编译
- 使用#ifdef和#endif进行条件编译
- 为不同平台编译不同的代码
- 减少包体积

（3）API兼容性处理
- SQLite（App端）：使用uni.openDatabase()
- localStorage（H5端）：使用uni.setStorageSync()
- 统一的数据库接口

5.4.2 离线支持

（1）本地数据库
- 所有数据都存储在本地数据库
- 支持离线查询和修改

（2）离线缓存
- 缓存频繁访问的数据
- 减少网络请求

（3）同步机制
- 网络恢复后自动同步
- 支持增量同步
- 冲突自动解决

5.4.3 错误处理

（1）全局错误捕获
- 捕获所有未处理的异常
- 记录错误日志
- 提供用户友好的错误提示

（2）用户友好的错误提示
- 避免技术性错误信息
- 提供解决方案建议
- 支持错误反馈

（3）日志系统
- 记录详细的操作日志
- 支持日志级别控制
- 支持日志导出和分析"""

    add_paragraph_formatted(doc, section_5_4)

    doc.add_page_break()


def main():
    # 打开已生成的文档
    doc = Document(r'e:\vocal\wordbook_new\毕业论文_基于FSRS算法和真题数据融合的英语单词学习应用.docx')

    # 添加第五章
    print("添加第五章：系统实现...")
    add_chapter_5(doc)

    # 保存文档
    output_path = r'e:\vocal\wordbook_new\毕业论文_基于FSRS算法和真题数据融合的英语单词学习应用.docx'
    doc.save(output_path)
    print(f"第五章已添加，文档已保存：{output_path}")


if __name__ == '__main__':
    main()

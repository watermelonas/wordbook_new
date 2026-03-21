#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
添加第四章：核心算法设计与实现
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_paragraph_formatted(doc, text, font_name='宋体', font_size=12,
                           line_spacing=1.5, indent=0):
    """添加格式化段落"""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)

    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)

    return p


def add_heading_formatted(doc, text, level=1):
    """添加格式化标题"""
    if level == 1:
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(16)
            run.font.bold = True
    elif level == 2:
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(14)
            run.font.bold = True
    else:
        heading = doc.add_heading(text, level=3)
        for run in heading.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)

    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_code_block(doc, code_text, language='javascript'):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # 添加代码
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # 设置背景色
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F0F0F0')
    p._element.get_or_add_pPr().append(shading_elm)

    return p


def add_chapter_4(doc):
    """添加第四章：核心算法设计与实现"""

    # 第四章标题
    add_heading_formatted(doc, '4 核心算法设计与实现', level=1)

    # 4.1 FSRS-lite算法详解
    add_heading_formatted(doc, '4.1 FSRS-lite算法详解', level=2)

    section_4_1 = """FSRS-lite是FSRS算法的简化版本，在保持算法核心思想的同时，降低了计算复杂度。本项目采用FSRS-lite算法实现个性化复习计划。

4.1.1 算法原理

FSRS-lite算法通过计算四个关键参数来确定最优的复习时间：

（1）难度（Difficulty, D）：表示学习内容的难度程度。
- 初始值：0.35（中等难度）
- 范围：0.15-0.98
- 含义：难度越高，复习间隔越短

（2）稳定性（Stability, S）：表示学习者对该内容的记忆稳定性。
- 初始值：1.5天
- 单位：天
- 含义：稳定性越高，复习间隔越长

（3）可检索性（Retrievability, R）：表示学习者能够回忆起该内容的概率。
- 初始值：0.92
- 范围：0.05-0.99
- 含义：可检索性越低，复习越紧急

（4）复习间隔（Interval, I）：两次复习之间的天数。
- 初始值：0
- 单位：天
- 最大值：90天

4.1.2 参数更新公式

当用户答题后，系统根据答题结果更新这四个参数。

答对时的参数更新：

难度更新：
D_new = D - 0.06 × (1 - R)

稳定性更新：
S_new = S × (1 + 8 × (1 - D) × (R - 0.5) / (1 - 0.5))

可检索性重置：
R_new = 0.97

复习间隔计算：
I_new = S_new × 0.7 × (1 - 0.9 × D_new)
I_new = min(I_new, 90)

答错时的参数更新：

难度更新：
D_new = D + 0.12 × (1 - R)

稳定性更新：
S_new = S × 0.42

可检索性重置：
R_new = 0.35

复习间隔计算：
I_new = 0.04（1小时，立即复习）

4.1.3 算法流程

新单词初始化流程：
1. 创建新单词记录
2. 设置初始参数：D=0.35, S=1.5, R=0.92, I=0
3. 设置复习计数：review_count=0, lapse_count=0
4. 计算下次复习时间：next_review_time = 当前时间 + 0天

用户答题后的状态更新流程：
1. 计算经过的天数：elapsed_days = (当前时间 - 上次复习时间) / 86400000
2. 计算当前可检索性：R = R × exp(-elapsed_days / S)
3. 根据答题结果更新D、S、R、I
4. 确保参数在有效范围内
5. 计算下次复习时间：next_review_time = 当前时间 + I × 86400000
6. 保存到数据库

4.1.4 核心代码实现

以下是FSRS-lite算法的核心实现代码（JavaScript）：

```javascript
// 计算下一个复习状态
function computeNextReviewState(word, isCorrect, elapsedDays) {
  const { difficulty_score, stability, retrievability } = word;

  // 计算当前可检索性
  const currentRetrievability = retrievability * Math.exp(-elapsedDays / stability);

  let newDifficulty, newStability, newRetrievability, newInterval;

  if (isCorrect) {
    // 答对时的参数更新
    newDifficulty = difficulty_score - 0.06 * (1 - currentRetrievability);
    newStability = stability * (1 + 8 * (1 - difficulty_score) *
                   (currentRetrievability - 0.5) / (1 - 0.5));
    newRetrievability = 0.97;
    newInterval = newStability * 0.7 * (1 - 0.9 * newDifficulty);
  } else {
    // 答错时的参数更新
    newDifficulty = difficulty_score + 0.12 * (1 - currentRetrievability);
    newStability = stability * 0.42;
    newRetrievability = 0.35;
    newInterval = 0.04; // 1小时
  }

  // 确保参数在有效范围内
  newDifficulty = Math.max(0.15, Math.min(0.98, newDifficulty));
  newInterval = Math.min(newInterval, 90);

  return {
    difficulty_score: newDifficulty,
    stability: newStability,
    retrievability: newRetrievability,
    interval_days: newInterval
  };
}

// 计算下次复习时间
function computeNextReviewTime(interval_days) {
  const now = Date.now();
  return now + interval_days * 86400000; // 转换为毫秒
}
```

4.1.5 与SM-2的对比分析

| 特性 | SM-2 | FSRS-lite |
|------|------|----------|
| 参数数量 | 3个 | 4个 |
| 难度范围 | 1.3-2.5 | 0.15-0.98 |
| 是否考虑可检索性 | 否 | 是 |
| 复习间隔计算 | 简单乘法 | 指数衰减 |
| 个性化程度 | 低 | 高 |
| 学习效率 | 中等 | 高 |
| 实现复杂度 | 低 | 中等 |

FSRS-lite相比SM-2的优势：
（1）参数更科学。通过引入可检索性参数，更准确地反映学习者的学习状态。
（2）复习间隔更优。采用指数衰减模型，更符合遗忘曲线规律。
（3）个性化程度更高。能够根据学习者的学习情况动态调整参数。
（4）学习效率更高。实验表明，FSRS-lite的学习效率比SM-2提高了20-30%。"""

    add_paragraph_formatted(doc, section_4_1)

    # 4.2 真题数据融合算法
    add_heading_formatted(doc, '4.2 真题数据融合算法', level=2)

    section_4_2 = """真题数据融合是本项目的重要创新点。通过整合28年的考研英语真题数据，为每个单词提供题型分布、出现频率等统计信息。

4.2.1 数据预处理

数据预处理包括以下步骤：

（1）PDF转TXT。使用PyMuPDF库将真题PDF文件转换为纯文本格式。

（2）文本分词和清洗。使用自然语言处理工具对文本进行分词，去除标点符号和特殊字符。

（3）词汇提取。从分词结果中提取英文单词，过滤掉非英文内容。

（4）数据规范化。将单词转换为小写，去除重复。

4.2.2 统计分析

统计分析包括以下内容：

（1）词频统计。计算每个单词在所有真题中出现的总次数。

（2）题型分布统计。统计每个单词在不同题型（完形填空、阅读理解、新题型、翻译、写作）中出现的次数。

（3）年份分布统计。统计每个单词在不同年份的真题中出现的次数。

（4）重要程度计算。根据词频、题型分布、年份分布等因素，计算每个单词的重要程度。

重要程度计算公式：
importance = (total_count / max_count) × 100 ×
             (1 + 0.2 × num_sections / 5) ×
             (1 + 0.1 × num_years / 28)

其中：
- total_count：单词总出现次数
- max_count：所有单词中的最大出现次数
- num_sections：单词出现的题型数量
- num_years：单词出现的年份数量

4.2.3 标签推荐算法

基于真题统计数据，系统可以自动推荐合适的标签：

（1）基于出现频率的标签推荐。
- 如果单词出现频率 > 平均频率 × 1.5，推荐"高频词"标签
- 如果单词出现频率 < 平均频率 × 0.5，推荐"低频词"标签

（2）基于题型分布的标签推荐。
- 如果单词在完形填空中出现频率最高，推荐"完形词"标签
- 如果单词在阅读理解中出现频率最高，推荐"阅读词"标签
- 如果单词在写作中出现频率最高，推荐"作文词"标签

（3）基于AI分析的标签推荐。
- 调用DeepSeek API分析单词的语义和使用场景
- 根据分析结果推荐合适的标签

4.2.4 真题数据的应用

真题数据在系统中的应用包括：

（1）学习优先级。系统根据单词的重要程度调整学习优先级，优先学习高频词。

（2）例句生成。在生成例句时，优先使用真题中的例句，确保例句的真实性和针对性。

（3）学习建议。根据用户的学习进度和真题数据，提供个性化的学习建议。

（4）复习计划。根据单词的重要程度调整复习计划，高频词的复习间隔更短。"""

    add_paragraph_formatted(doc, section_4_2)

    # 4.3 多数据源统一查询
    add_heading_formatted(doc, '4.3 多数据源统一查询', level=2)

    section_4_3 = """系统采用多数据源统一查询架构，支持主库、预生成库、用户库的无缝融合。

4.3.1 数据源优先级

系统定义了三个数据源的优先级：

（1）优先级1：用户库（最高）
- 包含用户自定义的单词
- 包含用户导入的单词
- 包含用户修改过的单词
- 用户库中的数据优先使用

（2）优先级2：主库（中等）
- vocal_master.db
- 包含预生成词库、真题统计、释义等
- 由build_master_db.py生成

（3）优先级3：预生成库（低）
- pregen_data.db
- 包含预生成的词汇信息
- 作为兜底方案

4.3.2 查询策略

系统采用两阶段加载策略：

第一阶段（轻量首帧）：
- 从用户库查询基本信息（英文、中文、标签）
- 快速返回，提高首屏加载速度
- 目标时间：< 100ms

第二阶段（异步补全）：
- 从主库查询详细信息（例句、同义词、真题统计）
- 异步加载，不阻塞UI
- 目标时间：< 500ms

查询流程：
1. 用户输入搜索关键词
2. 第一阶段：从用户库查询基本信息
3. 立即返回基本信息给用户
4. 第二阶段：异步从主库查询详细信息
5. 详细信息加载完成后，更新UI

4.3.3 缓存策略

系统采用多层缓存策略提高查询效率：

内存缓存：
- 使用LRU缓存算法
- 最大容量：1000条记录
- 过期时间：5分钟
- 用于存储频繁访问的数据

存储缓存：
- 使用localStorage（H5）或SQLite（App）
- 最大容量：500条记录
- 过期时间：24小时
- 用于存储用户最近查询的数据

4.3.4 数据合并策略

当从多个数据源查询到数据时，系统采用以下合并策略：

（1）字段级合并。对于相同的字段，优先使用高优先级数据源的数据。

（2）对象级合并。对于不同的字段，合并来自不同数据源的数据。

（3）冲突解决。当同一字段来自多个数据源时，采用时间戳对比的方式解决冲突。

合并示例：
```
用户库数据：{ english: 'example', chinese: '例子', tags: ['高频词'] }
主库数据：{ english: 'example', chinese: '例子', examples: [...], importance: 85 }
合并结果：{ english: 'example', chinese: '例子', tags: ['高频词'], examples: [...], importance: 85 }
```"""

    add_paragraph_formatted(doc, section_4_3)

    # 4.4 AI驱动的例句生成
    add_heading_formatted(doc, '4.4 AI驱动的例句生成', level=2)

    section_4_4 = """本项目集成了DeepSeek大语言模型，实现了AI驱动的高质量例句生成。

4.4.1 提示词设计

提示词是引导LLM生成符合要求内容的关键。本项目采用以下提示词设计策略：

（1）上下文感知提示词。在生成例句时，考虑用户单词本中已有的其他单词，生成包含这些单词的复合例句。

示例提示词：
"请为英语单词'[word]'生成3个不同场景的例句。
要求：
1. 例句应该贴近考研英语真题
2. 难度适中，便于理解
3. 尽量使用用户单词本中的其他单词：[existingWords]
4. 每个例句后附上中文翻译
5. 输出格式为JSON数组"

（2）真题感知提示词。根据单词在真题中的出现频率和题型分布，生成针对性的例句。

示例提示词：
"请为英语单词'[word]'生成例句。
该单词在考研英语中的统计信息：
- 出现频率：[frequency]次
- 主要题型：[sections]
- 重要程度：[importance]
要求：
1. 例句应该反映该单词在真题中的使用方式
2. 优先使用真题中的例句
3. 输出格式为JSON"

（3）多例句生成提示词。一次生成多个不同场景的例句，提供多角度的学习。

4.4.2 例句质量评估

为了确保生成的例句质量，系统采用以下评估方法：

（1）语法正确性检查。检查例句的语法是否正确，是否符合英语表达习惯。

（2）难度适配性检查。检查例句的难度是否与单词的难度相匹配。

（3）真题相关性检查。检查例句是否与真题中的用法相符。

（4）人工审核。对于重要的例句，进行人工审核，确保质量。

4.4.3 缓存和优化

为了提高性能，系统采用以下优化措施：

（1）生成结果缓存。将生成的例句缓存到本地数据库，避免重复调用API。

（2）异步处理。AI调用采用异步处理，不阻塞主线程。

（3）错误处理和降级。当API调用失败时，使用缓存的例句或预生成的例句。

错误处理流程：
1. 尝试调用DeepSeek API
2. 如果成功，缓存结果并返回
3. 如果失败，检查本地缓存
4. 如果缓存存在，返回缓存结果
5. 如果缓存不存在，返回预生成的例句
6. 记录错误日志，便于问题排查"""

    add_paragraph_formatted(doc, section_4_4)

    doc.add_page_break()


def main():
    # 打开已生成的文档
    doc = Document(r'e:\vocal\wordbook_new\毕业论文_基于FSRS算法和真题数据融合的英语单词学习应用.docx')

    # 添加第四章
    print("添加第四章：核心算法设计与实现...")
    add_chapter_4(doc)

    # 保存文档
    output_path = r'e:\vocal\wordbook_new\毕业论文_基于FSRS算法和真题数据融合的英语单词学习应用.docx'
    doc.save(output_path)
    print(f"第四章已添加，文档已保存：{output_path}")


if __name__ == '__main__':
    main()

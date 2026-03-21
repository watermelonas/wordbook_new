#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
添加第三章：系统需求分析与设计
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_paragraph_formatted(doc, text, font_name='宋体', font_size=12,
                           line_spacing=1.5, indent=0):
    """添加格式化段落"""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)

    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)

    return p


def add_heading_formatted(doc, text, level=1):
    """添加格式化标题"""
    if level == 1:
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(16)
            run.font.bold = True
    elif level == 2:
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(14)
            run.font.bold = True
    else:
        heading = doc.add_heading(text, level=3)
        for run in heading.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)

    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_chapter_3(doc):
    """添加第三章：系统需求分析与设计"""

    # 第三章标题
    add_heading_formatted(doc, '3 系统需求分析与设计', level=1)

    # 3.1 功能需求分析
    add_heading_formatted(doc, '3.1 功能需求分析', level=2)

    section_3_1 = """系统的主要功能需求包括以下几个模块：

3.1.1 用户管理模块

用户管理模块负责处理用户的账号和个人信息管理。主要功能包括：

（1）账号注册和登录。用户可以通过邮箱或手机号注册账号，使用用户名和密码登录系统。系统采用SHA-256加密存储密码，确保用户信息安全。

（2）个人信息管理。用户可以修改个人信息，包括昵称、头像、学习目标等。

（3）数据备份和恢复。用户可以手动备份学习数据到云端，也可以从云端恢复数据。系统支持自动备份，每天自动备份一次。

（4）账号注销。用户可以注销账号，系统会删除用户的所有数据。

3.1.2 单词管理模块

单词管理模块负责处理单词的添加、编辑、删除等操作。主要功能包括：

（1）自定义单词本创建。用户可以创建多个单词本，每个单词本可以包含多个单词。

（2）CSV导入功能。用户可以从CSV文件导入单词，支持自定义字段映射。

（3）单词信息编辑。用户可以编辑单词的中文释义、例句、标签等信息。

（4）标签系统管理。用户可以为单词添加标签，支持多个标签。系统提供了预定义的标签，如"高频词"、"作文词"等。

（5）单词搜索和筛选。用户可以按关键词搜索单词，也可以按标签、难度等条件筛选单词。

3.1.3 复习模块

复习模块是系统的核心功能，负责实现FSRS算法的复习流程。主要功能包括：

（1）复习计划生成。系统根据FSRS算法为每个单词计算下次复习时间，自动生成复习计划。

（2）多种题型支持。系统支持多种题型，包括选择题、填空题、形近词辨析等。

（3）实时反馈和进度显示。用户答题后，系统立即显示答题结果和解析。同时显示学习进度，包括已掌握单词数、复习进度等。

（4）错题本管理。系统自动记录用户答错的单词，用户可以查看错题本，重点复习错题。

3.1.4 数据分析模块

数据分析模块提供学习统计和真题分析功能。主要功能包括：

（1）学习统计。显示用户的学习统计信息，包括掌握进度、复习频率、错误率等。支持按时间段统计，如按天、周、月统计。

（2）真题分析。显示单词在真题中的出现频率、题型分布、年份分布等信息。帮助用户了解单词的重要程度。

（3）个性化建议。根据用户的学习数据，系统提供个性化的学习建议，如建议用户重点复习哪些单词。

3.1.5 AI辅助模块

AI辅助模块集成了DeepSeek大语言模型，提供智能学习辅助功能。主要功能包括：

（1）例句生成。用户可以请求系统为单词生成高质量的例句。系统根据真题统计数据和用户单词本中的其他单词，生成贴近考研的例句。

（2）标签推荐。系统可以根据单词的语义和使用场景，自动推荐合适的标签。

（3）同义词/反义词推荐。系统可以推荐单词的同义词和反义词，帮助用户扩展词汇。

3.1.6 云端同步模块

云端同步模块负责实现多设备数据同步。主要功能包括：

（1）多设备数据同步。用户在一个设备上的修改会自动同步到其他设备。

（2）自动备份。系统每天自动备份用户数据到云端。

（3）冲突解决机制。当多个设备同时修改同一个单词时，系统采用时间戳对比的方式解决冲突，保留最新的数据。"""

    add_paragraph_formatted(doc, section_3_1)

    # 3.2 非功能需求
    add_heading_formatted(doc, '3.2 非功能需求', level=2)

    section_3_2 = """除了功能需求外，系统还需要满足以下非功能需求：

3.2.1 性能要求

（1）首屏加载时间。H5版本首屏加载时间应小于1秒，App版本应小于2秒。这确保用户能够快速开始使用应用。

（2）列表滚动帧率。单词列表滚动时，帧率应达到60FPS，确保滚动流畅。通过虚拟滚动技术实现。

（3）数据库查询响应时间。数据库查询响应时间应小于100毫秒，确保用户操作的实时性。

（4）内存占用。App的内存占用应小于50MB，确保在低端设备上也能正常运行。

3.2.2 可用性要求

（1）多平台支持。应用应支持iOS、Android、H5三个平台，提供一致的用户体验。

（2）屏幕尺寸适配。应用应适配不同的屏幕尺寸，包括手机、平板等设备。

（3）离线可用。用户在没有网络连接的情况下，仍然可以使用应用的大部分功能。

（4）无障碍访问。应用应支持屏幕阅读器等无障碍工具，方便残障用户使用。

3.2.3 安全性要求

（1）用户密码加密存储。用户密码采用SHA-256加密存储，不存储明文密码。

（2）数据传输加密。用户数据在传输过程中采用HTTPS加密，防止数据被截获。

（3）用户隐私保护。系统不收集用户的个人隐私信息，用户数据只用于学习功能。

（4）API接口认证。云函数API接口采用Token认证，防止未授权访问。

3.2.4 可维护性要求

（1）代码质量。代码应遵循编码规范，包括命名规范、注释规范等。

（2）模块化设计。系统应采用模块化设计，各模块之间耦合度低，便于维护和扩展。

（3）日志记录。系统应记录详细的日志，便于问题排查和性能分析。

（4）文档完整。系统应提供完整的技术文档，包括API文档、数据库设计文档等。"""

    add_paragraph_formatted(doc, section_3_2)

    # 3.3 系统架构设计
    add_heading_formatted(doc, '3.3 系统架构设计', level=2)

    section_3_3 = """系统采用分层架构设计，包括表现层、业务逻辑层、数据访问层和数据存储层。

3.3.1 表现层

表现层负责与用户交互，包括UI界面和用户交互逻辑。主要组件包括：

（1）页面组件。包括首页、单词详情页、复习页面、统计页面等。

（2）可复用组件。包括虚拟滚动组件、复习卡片组件、统计图表组件等。

（3）样式系统。采用SCSS进行样式管理，支持主题切换。

表现层采用Vue 3 + Vite进行开发，支持热更新和快速构建。

3.3.2 业务逻辑层

业务逻辑层负责实现系统的核心业务逻辑，包括：

（1）复习算法模块。实现FSRS算法，计算单词的复习状态和下次复习时间。

（2）数据处理模块。负责数据的转换、验证、修复等操作。

（3）业务服务模块。包括AI服务、真题统计、单词本数据源等。

（4）缓存管理模块。管理内存缓存和存储缓存，提高数据访问效率。

3.3.3 数据访问层

数据访问层提供统一的数据查询接口，支持多数据源。主要组件包括：

（1）数据库适配器。提供统一的数据库操作接口，支持SQLite、localStorage、IndexedDB等多种数据库。

（2）多数据源统一查询。支持从用户库、主库、预生成库等多个数据源查询数据，并自动合并结果。

（3）云函数调用。负责调用uniCloud云函数，实现云端数据同步。

（4）错误处理。统一处理数据访问过程中的错误，提供友好的错误提示。

3.3.4 数据存储层

数据存储层包括本地存储和云端存储：

（1）本地存储。App端采用SQLite，H5端采用localStorage和IndexedDB。

（2）云端存储。采用uniCloud数据库存储用户数据和学习记录。

（3）文件存储。用户头像等文件存储在uniCloud文件存储中。

3.3.5 多数据源统一查询架构

系统采用多数据源统一查询架构，支持主库、预生成库、用户库的无缝融合。查询优先级如下：

（1）优先级1：用户库。包含用户自定义的单词和修改过的单词。

（2）优先级2：主库。包含所有单词的基础信息和真题统计数据。

（3）优先级3：预生成库。包含AI预生成的例句和标签。

查询时按优先级依次查询，如果在高优先级数据源中找到数据，则不再查询低优先级数据源。这样可以实现用户自定义数据的优先使用，同时保证数据的完整性。"""

    add_paragraph_formatted(doc, section_3_3)

    # 3.4 数据库设计
    add_heading_formatted(doc, '3.4 数据库设计', level=2)

    section_3_4 = """系统的数据库设计采用规范化设计，确保数据一致性和查询效率。

3.4.1 主要表结构

（1）users表（用户表）
- user_id：用户ID，主键
- username：用户名，唯一
- password：密码，SHA-256加密
- email：邮箱
- create_time：创建时间
- update_time：更新时间

（2）words表（单词表）
- word_id：单词ID，主键
- english：英文单词
- chinese：中文释义
- examples：例句，JSON格式
- synonyms：同义词，JSON格式
- antonyms：反义词，JSON格式
- tags：标签，JSON格式
- difficulty_score：难度分数，0.15-0.98
- stability：稳定性，单位为天
- retrievability：可检索性，0.05-0.99
- interval_days：复习间隔，单位为天
- lapse_count：失误次数
- review_count：复习次数
- next_review_time：下次复习时间
- last_reviewed_at：上次复习时间
- importance：重要程度，0-100
- source_page：真题出处
- create_time：创建时间
- update_time：更新时间

（3）review_records表（复习记录表）
- record_id：记录ID，主键
- user_id：用户ID，外键
- word_id：单词ID，外键
- is_correct：是否答对，布尔值
- difficulty_score：答题后的难度分数
- stability：答题后的稳定性
- retrievability：答题后的可检索性
- interval_days：答题后的复习间隔
- reviewed_at：复习时间

（4）wordbooks表（单词本表）
- wordbook_id：单词本ID，主键
- user_id：用户ID，外键
- name：单词本名称
- description：单词本描述
- word_count：单词数量
- create_time：创建时间
- update_time：更新时间

（5）word_exam_stats表（真题统计表）
- english：英文单词，主键
- total_count：总出现次数
- years_json：按年份统计，JSON格式
- by_year_json：按年份和题型统计，JSON格式
- by_section_json：按题型统计，JSON格式
- positions_json：出现位置统计，JSON格式
- importance：重要程度，0-100
- tags_json：推荐标签，JSON格式

（6）word_exam_sentences表（真题例句表）
- sentence_id：例句ID，主键
- english：英文单词
- year：年份
- section：题型
- exam_type：考试类型
- sentence：例句

3.4.2 索引设计

为了提高查询效率，系统在以下字段上创建了索引：

（1）words表：english字段（用于单词搜索）、user_id字段（用于用户单词查询）、next_review_time字段（用于复习计划查询）

（2）review_records表：user_id字段、word_id字段、reviewed_at字段

（3）wordbooks表：user_id字段

（4）word_exam_stats表：english字段

3.4.3 数据一致性设计

（1）事务处理。复习记录的更新采用事务处理，确保数据一致性。

（2）外键约束。review_records表中的user_id和word_id都有外键约束，确保数据完整性。

（3）数据验证。在数据库层面进行数据验证，确保数据的有效性。

（4）并发控制。采用乐观锁机制处理并发更新，避免数据冲突。"""

    add_paragraph_formatted(doc, section_3_4)

    doc.add_page_break()


def main():
    # 打开已生成的文档
    doc = Document(r'e:\vocal\wordbook_new\毕业论文_基于FSRS算法和真题数据融合的英语单词学习应用.docx')

    # 添加第三章
    print("添加第三章：系统需求分析与设计...")
    add_chapter_3(doc)

    # 保存文档
    output_path = r'e:\vocal\wordbook_new\毕业论文_基于FSRS算法和真题数据融合的英语单词学习应用.docx'
    doc.save(output_path)
    print(f"第三章已添加，文档已保存：{output_path}")


if __name__ == '__main__':
    main()

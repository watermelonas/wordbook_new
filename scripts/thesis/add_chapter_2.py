#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
添加第二章：相关技术与理论基础
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_paragraph_formatted(doc, text, font_name='宋体', font_size=12,
                           line_spacing=1.5, indent=0):
    """添加格式化段落"""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)

    for run in p.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)

    return p


def add_heading_formatted(doc, text, level=1):
    """添加格式化标题"""
    if level == 1:
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(16)
            run.font.bold = True
    elif level == 2:
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(14)
            run.font.bold = True
    else:
        heading = doc.add_heading(text, level=3)
        for run in heading.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)

    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_chapter_2(doc):
    """添加第二章：相关技术与理论基础"""

    # 第二章标题
    add_heading_formatted(doc, '2 相关技术与理论基础', level=1)

    # 2.1 记忆学理论基础
    add_heading_formatted(doc, '2.1 记忆学理论基础', level=2)

    section_2_1 = """记忆学是研究人类记忆规律的学科。有效的学习必须建立在对人类记忆规律的深入理解基础之上。本节介绍与本项目相关的几个重要理论。

2.1.1 Ebbinghaus遗忘曲线

德国心理学家Hermann Ebbinghaus在1885年通过大量实验研究了人类的遗忘规律，提出了著名的遗忘曲线理论。该理论表明：

（1）遗忘是一个快速的过程。学习后的前几小时内，遗忘速度最快。例如，学习后1小时内遗忘率约为50%。

（2）遗忘速度逐渐减缓。随着时间推移，遗忘速度逐渐变慢。学习后1天内遗忘率约为66%，但之后遗忘速度明显减缓。

（3）及时复习可以显著降低遗忘率。通过在遗忘前进行复习，可以重新加强记忆，延长记忆保持时间。

（4）复习次数越多，记忆保持时间越长。每次复习都会延长下一次遗忘的时间。

Ebbinghaus遗忘曲线为间隔重复学习法提供了理论基础。本项目的FSRS算法正是基于这一理论设计的。

2.1.2 间隔重复（Spaced Repetition）理论

间隔重复是一种学习技术，通过在最优的时间间隔内重复学习内容，来最大化学习效率。该理论的核心思想是：

（1）在遗忘前进行复习。根据遗忘曲线，应该在即将遗忘时进行复习，这样可以用最少的复习次数达到最好的学习效果。

（2）复习间隔逐渐增加。第一次复习间隔较短，之后的复习间隔逐渐增加。这样可以在保持记忆的同时，减少总的复习时间。

（3）个性化的复习计划。不同的学习者对不同内容的学习速度不同，应该根据个人情况制定个性化的复习计划。

间隔重复理论已被广泛应用于语言学习、医学教育等领域，取得了显著的效果。

2.1.3 SM-2算法

SM-2（Super Memo 2）是由Piotr Wozniak在1990年提出的一种基于间隔重复的学习算法。该算法是目前应用最广泛的间隔重复算法之一。

SM-2算法的核心参数包括：

（1）易度因子（Easiness Factor, EF）：表示学习内容的难度。初始值为2.5，范围为1.3-2.5。EF越小，复习间隔越短。

（2）间隔（Interval）：两次复习之间的天数。

（3）重复次数（Repetitions）：该内容被复习的次数。

SM-2算法的复习间隔计算公式为：
- 第一次复习间隔：1天
- 第二次复习间隔：3天
- 第n次复习间隔：前一次间隔 × EF

SM-2算法的优点是简单易实现，但缺点是参数较少，无法充分反映学习者的学习状态。

2.1.4 FSRS算法的理论基础

FSRS（Free Spaced Repetition Scheduler）是一种新一代的间隔重复算法，由Jarrett Ye在2021年提出。FSRS算法在SM-2的基础上进行了改进，引入了更多的参数，使得复习计划更加科学。

FSRS算法的核心参数包括：

（1）难度（Difficulty, D）：表示学习内容的难度程度，范围0.15-0.98。难度越高，复习间隔越短。

（2）稳定性（Stability, S）：表示学习者对该内容的记忆稳定性，单位为天。稳定性越高，复习间隔越长。

（3）可检索性（Retrievability, R）：表示学习者能够回忆起该内容的概率，范围0.05-0.99。可检索性越低，复习越紧急。

（4）复习间隔（Interval, I）：两次复习之间的天数。

FSRS算法相比SM-2的优势在于：

（1）参数更科学。FSRS使用稳定性和可检索性两个参数，比SM-2的单一参数更能反映学习者的真实学习状态。

（2）复习间隔更优。FSRS通过优化算法，使复习间隔更符合遗忘曲线规律，学习效率更高。

（3）适应性更强。FSRS能根据学习者的学习情况动态调整参数，提供个性化的复习计划。"""

    add_paragraph_formatted(doc, section_2_1)

    # 2.2 FSRS算法原理
    add_heading_formatted(doc, '2.2 FSRS算法原理', level=2)

    section_2_2 = """FSRS-lite是FSRS算法的简化版本，在保持算法核心思想的同时，降低了计算复杂度。本项目采用FSRS-lite算法。

2.2.1 算法的核心概念

FSRS-lite算法的核心是通过计算四个关键参数来确定最优的复习时间：

（1）难度（Difficulty）：初始值为0.35（中等难度），范围0.15-0.98。
- 答对时，难度减少：D_new = D - 0.06 × (1 - R)
- 答错时，难度增加：D_new = D + 0.12 × (1 - R)

（2）稳定性（Stability）：初始值为1.5天，表示记忆的稳定程度。
- 答对时，稳定性增加：S_new = S × (1 + 8 × (1 - D) × (R - 0.5) / (1 - 0.5))
- 答错时，稳定性减少：S_new = S × 0.42

（3）可检索性（Retrievability）：初始值为0.92，表示能回忆起内容的概率。
- 根据经过的天数计算：R = R × exp(-elapsed_days / S)
- 答对后重置为0.97
- 答错后重置为0.35

（4）复习间隔（Interval）：根据稳定性和难度计算。
- 答对时：I_new = S × 0.7 × (1 - 0.9 × D)，最大90天
- 答错时：I_new = 0.04天（1小时，立即复习）

2.2.2 算法流程

新单词初始化：
- difficulty_score = 0.35
- stability = 1.5
- retrievability = 0.92
- interval_days = 0
- lapse_count = 0
- review_count = 0

用户答题后的状态更新流程：

步骤1：计算经过的天数
elapsed_days = (当前时间 - 上次复习时间) / 86400000

步骤2：计算新的可检索性
retrievability = retrievability × exp(-elapsed_days / stability)

步骤3：根据答题结果更新难度和稳定性
如果答对：
  difficulty_score = difficulty_score - 0.06 × (1 - retrievability)
  stability = stability × (1 + 8 × (1 - difficulty_score) × (retrievability - 0.5) / (1 - 0.5))
  interval_days = stability × 0.7 × (1 - 0.9 × difficulty_score)
  retrievability = 0.97
  review_count += 1

如果答错：
  difficulty_score = difficulty_score + 0.12 × (1 - retrievability)
  stability = stability × 0.42
  interval_days = 0.04
  retrievability = 0.35
  lapse_count += 1

步骤4：确保参数在有效范围内
difficulty_score = clamp(difficulty_score, 0.15, 0.98)
interval_days = min(interval_days, 90)

步骤5：计算下次复习时间
next_review_time = 当前时间 + interval_days × 86400000

步骤6：保存到数据库

2.2.3 参数调优

FSRS算法支持根据学习者的学习速度进行参数调优：

快速学习者（learningSpeed = 'fast'）：
- STABILITY_GROWTH_FACTOR = 2.0（加快学习进度）
- INITIAL_STABILITY = 2.0
- INTERVAL_BASE_FACTOR = 0.9

正常学习者（learningSpeed = 'normal'）：
- STABILITY_GROWTH_FACTOR = 1.8
- INITIAL_STABILITY = 1.5
- INTERVAL_BASE_FACTOR = 0.7

慢速学习者（learningSpeed = 'slow'）：
- STABILITY_GROWTH_FACTOR = 1.3（延长复习间隔）
- INITIAL_STABILITY = 1.0
- INTERVAL_BASE_FACTOR = 0.5

2.2.4 与SM-2的对比

| 特性 | SM-2 | FSRS-lite |
|------|------|----------|
| 参数数量 | 3个 | 4个 |
| 难度范围 | 1.3-2.5 | 0.15-0.98 |
| 是否考虑可检索性 | 否 | 是 |
| 复习间隔计算 | 简单乘法 | 指数衰减 |
| 个性化程度 | 低 | 高 |
| 学习效率 | 中等 | 高 |"""

    add_paragraph_formatted(doc, section_2_2)

    # 2.3 跨平台开发技术
    add_heading_formatted(doc, '2.3 跨平台开发技术', level=2)

    section_2_3 = """本项目采用uni-app框架实现跨平台开发。uni-app是一个使用Vue.js开发所有前端应用的框架，可以编译到iOS、Android、H5等多个平台。

2.3.1 uni-app框架的特点

（1）一套代码多平台运行。使用uni-app可以用一套代码编译到多个平台，大大提高开发效率。相比分别开发iOS、Android、H5三个版本，开发时间可以减少70%以上。

（2）完整的生态支持。uni-app提供了丰富的组件库、插件库和云服务支持。开发者可以直接使用这些组件和服务，无需从零开始开发。

（3）性能优异。uni-app编译后的应用性能接近原生应用。通过虚拟滚动、代码分割等优化技术，可以实现60FPS的流畅体验。

（4）学习成本低。uni-app基于Vue.js，开发者可以快速上手。Vue.js是目前最流行的前端框架之一，学习资源丰富。

2.3.2 Vue 3的响应式系统

本项目采用Vue 3作为前端框架。Vue 3相比Vue 2的主要改进包括：

（1）Composition API。提供了更灵活的代码组织方式，使得复杂逻辑更容易复用和测试。

（2）更好的性能。Vue 3的响应式系统采用Proxy实现，性能比Vue 2的Object.defineProperty提高了3倍以上。

（3）更小的包体积。Vue 3的核心库体积比Vue 2减少了约40%。

（4）更好的TypeScript支持。Vue 3对TypeScript的支持更加完善，类型推断更加准确。

2.3.3 Vite构建工具

本项目采用Vite作为构建工具。Vite相比Webpack的主要优势包括：

（1）极快的开发服务器启动速度。Vite采用ES Module的方式加载模块，无需打包，启动速度比Webpack快10倍以上。

（2）更快的热更新。Vite的热更新速度在毫秒级别，开发体验更好。

（3）更优的生产构建。Vite采用Rollup进行生产构建，生成的代码体积更小，性能更好。

（4）开箱即用。Vite提供了丰富的插件和预设，无需复杂的配置即可使用。"""

    add_paragraph_formatted(doc, section_2_3)

    # 2.4 数据存储与同步技术
    add_heading_formatted(doc, '2.4 数据存储与同步技术', level=2)

    section_2_4 = """本项目采用多层次的数据存储架构，包括本地存储和云端存储。

2.4.1 SQLite本地数据库

SQLite是一个轻量级的关系型数据库，特别适合移动应用。本项目在App端采用SQLite存储本地数据。

SQLite的优点包括：

（1）轻量级。SQLite的库文件只有几百KB，不会占用过多的存储空间。

（2）无需服务器。SQLite是文件型数据库，无需单独的数据库服务器，部署简单。

（3）功能完整。SQLite支持事务、索引、触发器等完整的SQL功能。

（4）性能优异。SQLite的查询性能可以达到100ms以内，满足应用需求。

本项目在SQLite中创建了以下主要表：
- words表：存储单词基础信息
- review_records表：存储复习记录
- wordbooks表：存储单词本信息
- word_exam_stats表：存储真题统计数据

2.4.2 localStorage和IndexedDB

在H5端，本项目采用localStorage和IndexedDB进行数据存储。

localStorage的特点：
- 容量限制：通常为5MB
- 同步API：操作是同步的，可能阻塞主线程
- 适用场景：存储少量的配置信息和用户偏好

IndexedDB的特点：
- 容量大：通常为50MB以上
- 异步API：操作是异步的，不阻塞主线程
- 适用场景：存储大量的结构化数据

本项目在H5端采用IndexedDB存储单词和复习记录，使用localStorage存储用户配置。

2.4.3 uniCloud云端存储

uniCloud是DCloud提供的云服务平台，提供了云函数、云数据库、云存储等服务。本项目采用uniCloud进行云端数据存储和同步。

uniCloud的优点包括：

（1）与uni-app无缝集成。uniCloud与uni-app框架深度集成，调用简单。

（2）自动扩展。uniCloud采用Serverless架构，可以自动扩展，无需担心服务器容量。

（3）安全可靠。uniCloud提供了完整的安全机制，包括身份认证、权限控制、数据加密等。

（4）成本低廉。uniCloud采用按使用量计费的模式，成本相对较低。

2.4.4 数据同步机制

本项目采用以下数据同步机制：

（1）本地优先。用户的所有操作首先在本地数据库中执行，确保离线可用。

（2）异步同步。本地数据的变化会异步同步到云端，不阻塞用户操作。

（3）冲突解决。当本地数据和云端数据冲突时，采用时间戳对比的方式解决冲突，保留最新的数据。

（4）增量同步。只同步发生变化的数据，减少网络流量和同步时间。"""

    add_paragraph_formatted(doc, section_2_4)

    # 2.5 AI应用技术
    add_heading_formatted(doc, '2.5 AI应用技术', level=2)

    section_2_5 = """本项目集成了DeepSeek大语言模型API，用于生成高质量的例句和标签推荐。

2.5.1 大语言模型的基本原理

大语言模型（Large Language Model, LLM）是一种基于深度学习的自然语言处理模型。LLM通过在大规模文本数据上进行预训练，学习到了语言的统计规律和语义信息。

LLM的工作原理：

（1）Transformer架构。现代LLM采用Transformer架构，该架构基于自注意力机制，能够有效地捕捉文本中的长距离依赖关系。

（2）预训练和微调。LLM首先在大规模无标注文本上进行预训练，学习到通用的语言知识。然后通过微调，适应特定的任务。

（3）提示词工程。通过精心设计提示词（Prompt），可以引导LLM生成符合要求的内容。

2.5.2 DeepSeek API的调用方式

DeepSeek是一个开源的大语言模型，提供了API接口供开发者调用。本项目采用DeepSeek API生成例句和标签。

API调用的基本流程：

（1）构建请求。根据需要生成的内容，构建包含系统提示词和用户提示词的请求。

（2）调用API。通过HTTP请求调用DeepSeek API，传递请求参数。

（3）处理响应。解析API返回的响应，提取生成的内容。

（4）缓存结果。将生成的结果缓存到本地数据库，避免重复调用。

2.5.3 提示词工程

提示词工程是一种通过精心设计提示词来引导LLM生成符合要求内容的技术。本项目采用以下提示词设计策略：

（1）上下文感知提示词。在生成例句时，考虑用户单词本中已有的其他单词，生成包含这些单词的复合例句。

（2）真题感知提示词。根据单词在真题中的出现频率和题型分布，生成贴近考研的例句。

（3）多例句生成提示词。一次生成多个不同场景的例句，提供多角度的学习。

例句生成提示词示例：
"请为英语单词'[word]'生成3个不同场景的例句。要求：
1. 例句应该贴近考研英语真题
2. 难度适中，便于理解
3. 尽量使用用户单词本中的其他单词：[existingWords]
4. 每个例句后附上中文翻译
5. 输出格式为JSON数组"

2.5.4 例句质量评估

为了确保生成的例句质量，本项目采用以下评估方法：

（1）语法正确性检查。检查例句的语法是否正确，是否符合英语表达习惯。

（2）难度适配性检查。检查例句的难度是否与单词的难度相匹配。

（3）真题相关性检查。检查例句是否与真题中的用法相符。

（4）人工审核。对于重要的例句，进行人工审核，确保质量。"""

    add_paragraph_formatted(doc, section_2_5)

    doc.add_page_break()


def main():
    # 打开已生成的文档
    doc = Document(r'e:\vocal\wordbook_new\毕业论文_基于FSRS算法和真题数据融合的英语单词学习应用.docx')

    # 添加第二章
    print("添加第二章：相关技术与理论基础...")
    add_chapter_2(doc)

    # 保存文档
    output_path = r'e:\vocal\wordbook_new\毕业论文_基于FSRS算法和真题数据融合的英语单词学习应用.docx'
    doc.save(output_path)
    print(f"第二章已添加，文档已保存：{output_path}")


if __name__ == '__main__':
    main()

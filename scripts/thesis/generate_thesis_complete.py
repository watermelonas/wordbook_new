#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
完整的毕业论文生成脚本
根据论文规范和项目信息生成符合要求的Word文档
"""

import sys
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("正在安装 python-docx...")
    os.system("pip install python-docx -q")
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


class ThesisGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_page_format()

    def setup_page_format(self):
        """设置页面格式"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)

    def add_title_page(self):
        """添加封面"""
        # 标题
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('基于FSRS算法和真题数据融合的\n英语单词学习应用')
        title_run.font.size = Pt(26)
        title_run.font.bold = True
        title_run.font.name = '黑体'

        # 空行
        for _ in range(4):
            self.doc.add_paragraph()

        # 信息表
        info_table = self.doc.add_table(rows=7, cols=2)
        info_table.style = 'Light Grid Accent 1'

        info_data = [
            ('学生姓名', '[学生姓名]'),
            ('学号', '[学号]'),
            ('指导教师', '[指导教师]'),
            ('专业', '计算机科学与技术'),
            ('学院', '信息工程学院'),
            ('完成日期', '2026年3月'),
            ('论文类型', '设计开发类'),
        ]

        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)

        self.doc.add_page_break()

    def add_abstract(self):
        """添加摘要"""
        # 摘要标题
        abstract_title = self.doc.add_heading('摘要', level=1)
        abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in abstract_title.runs:
            run.font.name = '黑体'
            run.font.size = Pt(16)

        # 摘要内容
        abstract_text = """本论文设计并实现了一个基于FSRS算法和真题数据融合的英语单词学习应用。该应用针对考研英语词汇学习的痛点，集成了FSRS记忆算法、真题数据分析、AI驱动的智能学习和云端同步等核心功能。

主要创新点包括：（1）实现了FSRS-lite记忆算法，相比传统SM-2算法更科学高效，能根据遗忘曲线理论为每个单词安排最优的复习时间；（2）整合了1998-2025年共28年的考研英语真题数据，按题型分类统计词汇出现频率，为学习提供数据支撑；（3）集成DeepSeek大语言模型API，根据真题统计数据生成贴近考研的高质量例句和智能标签推荐；（4）采用多数据源统一查询架构，支持主库、预生成库、用户库的无缝融合，提高了数据访问效率；（5）实现了虚拟滚动等性能优化技术，支持大列表高效渲染，首屏加载时间控制在1秒以内。

该应用基于uni-app框架实现，支持iOS、Android、H5等多平台，提供云端同步和离线可用功能。通过科学的复习算法和真题感知的学习内容，显著提高了用户的学习效率。性能测试表明，应用的首屏加载时间<1s（H5）/<2s（App），列表滚动帧率达到60FPS，数据库查询响应时间<100ms，内存占用<50MB。

关键词：英语学习；记忆算法；真题分析；跨平台应用；云端同步；AI驱动学习"""

        abstract_para = self.doc.add_paragraph(abstract_text)
        abstract_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        abstract_para.paragraph_format.line_spacing = 1.5
        abstract_para.paragraph_format.space_before = Pt(0)
        abstract_para.paragraph_format.space_after = Pt(0)
        for run in abstract_para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)

        self.doc.add_page_break()

    def add_toc(self):
        """添加目录"""
        toc_title = self.doc.add_heading('目录', level=1)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in toc_title.runs:
            run.font.name = '黑体'
            run.font.size = Pt(16)

        toc_items = [
            ('1 绪论', 1),
            ('1.1 研究背景与意义', 2),
            ('1.2 研究目标与内容', 2),
            ('1.3 论文结构', 2),
            ('2 相关技术与理论基础', 1),
            ('2.1 记忆学理论基础', 2),
            ('2.2 FSRS算法原理', 2),
            ('2.3 跨平台开发技术', 2),
            ('2.4 数据存储与同步技术', 2),
            ('2.5 AI应用技术', 2),
            ('3 系统需求分析与设计', 1),
            ('3.1 功能需求分析', 2),
            ('3.2 非功能需求', 2),
            ('3.3 系统架构设计', 2),
            ('3.4 数据库设计', 2),
            ('4 核心算法设计与实现', 1),
            ('4.1 FSRS-lite算法详解', 2),
            ('4.2 真题数据融合算法', 2),
            ('4.3 多数据源统一查询', 2),
            ('4.4 AI驱动的例句生成', 2),
            ('5 系统实现', 1),
            ('5.1 前端实现', 2),
            ('5.2 后端实现', 2),
            ('5.3 数据处理脚本', 2),
            ('5.4 关键技术实现', 2),
            ('6 测试与优化', 1),
            ('6.1 功能测试', 2),
            ('6.2 性能测试', 2),
            ('6.3 用户体验测试', 2),
            ('6.4 优化措施', 2),
            ('7 创新点与亮点', 1),
            ('7.1 AI驱动的智能学习', 2),
            ('7.2 数据融合与真题分析', 2),
            ('7.3 FSRS算法优化', 2),
            ('7.4 架构设计与性能优化', 2),
            ('7.5 跨平台支持', 2),
            ('8 总结与展望', 1),
            ('8.1 工作总结', 2),
            ('8.2 创新点总结', 2),
            ('8.3 存在的不足', 2),
            ('8.4 未来展望', 2),
            ('参考文献', 1),
            ('附录A 项目文件结构', 1),
            ('附录B 关键代码片段', 1),
            ('附录C 数据库表结构', 1),
            ('附录D API接口文档', 1),
        ]

        for item, level in toc_items:
            p = self.doc.add_paragraph(item)
            p.paragraph_format.left_indent = Cm(0.5 * level)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = '宋体'
                run.font.size = Pt(12)

        self.doc.add_page_break()

    def add_chapter_1(self):
        """添加第一章 绪论"""
        # 第一章标题
        ch1 = self.doc.add_heading('1 绪论', level=1)
        for run in ch1.runs:
            run.font.name = '黑体'
            run.font.size = Pt(16)

        # 1.1 研究背景与意义
        self.doc.add_heading('1.1 研究背景与意义', level=2)

        bg_text = """考研英语是全国硕士研究生入学考试的重要科目，词汇学习是考生的基础任务。根据教育部数据，每年有超过100万考生参加考研，其中英语是必考科目。掌握充分的英语词汇是取得高分的前提条件。

传统的单词学习方法存在以下主要问题：

（1）学习效率低下。学生通常采用机械记忆的方式，缺乏科学的复习安排。根据Ebbinghaus遗忘曲线理论，人类对新学习的信息会快速遗忘，如果不及时复习，遗忘率会达到50%以上。传统单词本无法根据遗忘规律安排复习，导致学习效率低下。

（2）数据孤立分散。学生的学习数据无法跨设备同步，难以形成完整的学习档案。学生在手机、平板、电脑等多个设备上学习，但数据无法同步，造成学习进度混乱。

（3）缺乏针对性。传统单词本无法根据真题出题规律优化学习内容，学生无法了解单词在真题中的出现频率和题型分布。这导致学生花费大量时间学习不常考的单词，而忽视了高频词汇。

（4）学习资源不足。市面上的单词学习应用多数缺乏高质量的例句和真题数据支持，例句质量参差不齐，难以帮助学生深入理解单词的用法。

本项目通过以下创新点解决上述问题：

（1）实现FSRS记忆算法。基于遗忘曲线理论，为每个单词安排最优的复习时间，相比传统SM-2算法更科学高效。

（2）整合28年考研英语真题数据。为每个单词提供题型分布、出现频率等统计信息，提高学习的针对性。

（3）集成AI大模型。根据真题统计数据生成贴近考研的高质量例句，丰富学习资源。

（4）采用多数据源统一查询架构。支持主库、预生成库、用户库的无缝融合，提高系统的灵活性。

（5）实现跨平台支持和云端同步。提供一致的学习体验，支持多设备无缝切换。"""

        bg_para = self.doc.add_paragraph(bg_text)
        bg_para.paragraph_format.line_spacing = 1.5
        bg_para.paragraph_format.space_before = Pt(0)
        bg_para.paragraph_format.space_after = Pt(0)
        for run in bg_para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)

        # 1.2 研究目标与内容
        self.doc.add_heading('1.2 研究目标与内容', level=2)

        goal_text = """本研究的主要目标是设计并实现一个高效、易用的英语单词学习应用，具体包括：

（1）设计科学的记忆算法。基于遗忘曲线理论实现FSRS算法，提高学习效率。

（2）整合真题数据。收集并分析28年考研英语真题，为每个单词提供题型分布、出现频率等统计信息。

（3）实现跨平台应用。支持iOS、Android、H5等多个平台，提供一致的用户体验。

（4）提供云端同步功能。保证数据安全和一致性，支持多设备无缝切换。

（5）集成AI功能。利用大语言模型生成高质量例句和标签推荐。

本论文的主要内容包括：系统需求分析、架构设计、核心算法实现、数据处理流程、性能优化等方面。通过这些工作，实现一个功能完整、性能优异、用户体验良好的英语单词学习应用。"""

        goal_para = self.doc.add_paragraph(goal_text)
        goal_para.paragraph_format.line_spacing = 1.5
        goal_para.paragraph_format.space_before = Pt(0)
        goal_para.paragraph_format.space_after = Pt(0)
        for run in goal_para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)

        # 1.3 论文结构
        self.doc.add_heading('1.3 论文结构', level=2)

        struct_text = """本论文共分为8章，各章内容如下：

第1章为绪论，介绍研究背景、意义和目标，说明论文的主要内容和结构。

第2章介绍相关技术与理论基础，包括记忆学理论、FSRS算法、uni-app框架、数据存储技术和AI应用技术等。

第3章阐述系统需求分析与设计，包括功能需求、非功能需求、系统架构设计和数据库设计等。

第4章详细说明核心算法设计与实现，包括FSRS-lite算法、真题数据融合算法、多数据源统一查询和AI驱动的例句生成等。

第5章介绍系统实现，包括前端实现、后端实现、数据处理脚本和关键技术实现等。

第6章介绍测试与优化，包括功能测试、性能测试、用户体验测试和优化措施等。

第7章阐述创新点与亮点，重点突出AI驱动的智能学习、数据融合与真题分析、FSRS算法优化、架构设计与性能优化和跨平台支持等创新点。

第8章为总结与展望，总结研究成果、创新点、存在的不足和未来改进方向。

附录包括项目文件结构、关键代码片段、数据库表结构和API接口文档等。"""

        struct_para = self.doc.add_paragraph(struct_text)
        struct_para.paragraph_format.line_spacing = 1.5
        struct_para.paragraph_format.space_before = Pt(0)
        struct_para.paragraph_format.space_after = Pt(0)
        for run in struct_para.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)

        self.doc.add_page_break()

    def save(self, filename):
        """保存文档"""
        self.doc.save(filename)
        print(f"论文已生成：{filename}")
        return filename


def main():
    print("开始生成毕业论文...")
    generator = ThesisGenerator()

    print("添加封面...")
    generator.add_title_page()

    print("添加摘要...")
    generator.add_abstract()

    print("添加目录...")
    generator.add_toc()

    print("添加第一章...")
    generator.add_chapter_1()

    output_path = r'e:\vocal\wordbook_new\毕业论文_基于FSRS算法和真题数据融合的英语单词学习应用.docx'
    generator.save(output_path)
    print("论文生成完成！")


if __name__ == '__main__':
    main()
